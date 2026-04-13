from uuid import uuid4

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.models.entities import Artifact
from app.services.artifacts import ArtifactService


def _make_artifact(*, title: str, storage_key: str, mime_type: str, kind: str = "document_export") -> Artifact:
    return Artifact(
        workspace_id=uuid4(),
        kind=kind,
        title=title,
        storage_key=storage_key,
        metadata={"mime_type": mime_type},
    )


def test_build_preview_for_csv_returns_table_data(tmp_path):
    service = ArtifactService()
    service.storage.root = tmp_path
    tmp_path.mkdir(parents=True, exist_ok=True)

    storage_key = "exports/test/metrics.csv"
    path = service.storage.resolve(storage_key, create_parent=True)
    path.write_text("region,revenue\nEU,10\nUS,20\n", encoding="utf-8")

    artifact = _make_artifact(
        title="metrics.csv",
        storage_key=storage_key,
        mime_type="text/csv",
    )

    preview = service.build_preview(artifact)

    assert preview["preview_kind"] == "csv"
    assert preview["table"]["columns"] == ["region", "revenue"]
    assert preview["table"]["rows"] == [["EU", "10"], ["US", "20"]]


def test_build_preview_for_spreadsheet_returns_sheet_preview(tmp_path):
    service = ArtifactService()
    service.storage.root = tmp_path
    tmp_path.mkdir(parents=True, exist_ok=True)

    storage_key = "exports/test/forecast.xlsx"
    path = service.storage.resolve(storage_key, create_parent=True)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Forecast"
    sheet.append(["Month", "Revenue"])
    sheet.append(["January", 120])
    sheet.append(["February", 140])
    workbook.save(path)
    workbook.close()

    artifact = _make_artifact(
        title="forecast.xlsx",
        storage_key=storage_key,
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    preview = service.build_preview(artifact)

    assert preview["preview_kind"] == "spreadsheet"
    assert preview["sheets"][0]["name"] == "Forecast"
    assert preview["sheets"][0]["columns"] == ["Month", "Revenue"]
    assert preview["sheets"][0]["rows"][0] == ["January", "120"]


def test_build_preview_for_docx_returns_text_content(tmp_path):
    service = ArtifactService()
    service.storage.root = tmp_path
    tmp_path.mkdir(parents=True, exist_ok=True)

    storage_key = "exports/test/brief.docx"
    path = service.storage.resolve(storage_key, create_parent=True)
    document = DocxDocument()
    document.add_heading("Launch Brief", level=1)
    document.add_paragraph("Coordinate the browser, analysis, and delivery agents.")
    document.save(path)

    artifact = _make_artifact(
        title="brief.docx",
        storage_key=storage_key,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    preview = service.build_preview(artifact)

    assert preview["preview_kind"] == "document"
    assert "Launch Brief" in preview["text_content"]
    assert "delivery agents" in preview["text_content"]
